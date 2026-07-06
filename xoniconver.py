from flask import Flask, request, make_response, render_template, flash, redirect, url_for
from PIL import Image
try:
    from pypdf import PdfReader, PdfWriter, PdfMerger
except ImportError:
    from PyPDF2 import PdfReader, PdfWriter, PdfMerger  # fallback legacy
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pdfplumber
import io
import zipfile
import datetime
import socket
import os
import re
import tempfile
import shutil
from werkzeug.utils import secure_filename
from io import StringIO
try:
    import qrcode
    _HAS_QRCODE = True
except ImportError:
    _HAS_QRCODE = False

app = Flask(__name__)
app.secret_key = "clave_secreta_xoni_conver_pc"

# Sin límite de tamaño
app.config['MAX_CONTENT_LENGTH'] = None

ALLOWED_IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "bmp", "gif", "tiff", "webp"}
ALLOWED_PDF_EXTENSIONS = {"pdf"}
ALLOWED_EXTENSIONS = ALLOWED_IMAGE_EXTENSIONS.union(ALLOWED_PDF_EXTENSIONS)

def generar_qr_terminal(texto):
    if not _HAS_QRCODE:
        return f"(instala qrcode para ver QR): {texto}"
    qr = qrcode.QRCode(version=1, box_size=2, border=1)
    qr.add_data(texto)
    qr.make(fit=True)
    qr_ascii = StringIO()
    qr.print_ascii(out=qr_ascii, invert=True)
    return qr_ascii.getvalue()

def allowed_filename(filename: str) -> bool:
    if not filename or '.' not in filename:
        return False
    ext = filename.rsplit(".", 1)[1].lower()
    return ext in ALLOWED_EXTENSIONS

def is_image_filename(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_IMAGE_EXTENSIONS

def is_pdf_filename(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_PDF_EXTENSIONS

# ────────────────────────────────────────────────
# Imágenes → PDF
# ────────────────────────────────────────────────
def images_to_pdf(files):
    pil_images = []
    temp_dirs = []
    try:
        for f in files:
            if not is_image_filename(f.filename):
                continue
            temp_dir = tempfile.mkdtemp()
            temp_path = os.path.join(temp_dir, secure_filename(f.filename))
            f.save(temp_path)
            temp_dirs.append(temp_dir)
            try:
                img = Image.open(temp_path).convert("RGB")
                pil_images.append(img)
            except Exception as e:
                print(f"Error procesando imagen {f.filename}: {e}")
                continue
        if not pil_images:
            raise ValueError("No hay imágenes válidas para convertir.")
        output = io.BytesIO()
        if len(pil_images) == 1:
            pil_images[0].save(output, format="PDF")
        else:
            first, rest = pil_images[0], pil_images[1:]
            first.save(output, format="PDF", save_all=True, append_images=rest)
        output.seek(0)
        return ("single", ("imagenes_a_pdf.pdf", output.read(), "application/pdf"))
    finally:
        for temp_dir in temp_dirs:
            try:
                shutil.rmtree(temp_dir)
            except:
                pass

# ────────────────────────────────────────────────
# Unir PDFs (método 1: PdfMerger)
# ────────────────────────────────────────────────
def pdfs_merge_v1(files):
    merger = PdfMerger()
    temp_dirs = []
    try:
        for idx, f in enumerate(files):
            if not is_pdf_filename(f.filename):
                continue
            temp_dir = tempfile.mkdtemp()
            temp_path = os.path.join(temp_dir, secure_filename(f.filename))
            f.save(temp_path)
            temp_dirs.append(temp_dir)
            merger.append(temp_path)
            print(f"✅ PDF {idx+1} agregado: {f.filename}")
        if len(merger.pages) == 0:
            raise ValueError("No se pudieron procesar archivos PDF válidos.")
        output = io.BytesIO()
        merger.write(output)
        merger.close()
        output.seek(0)
        return ("single", ("pdfs_combinados.pdf", output.read(), "application/pdf"))
    finally:
        for temp_dir in temp_dirs:
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except:
                pass

# ────────────────────────────────────────────────
# Unir PDFs (método 2: PdfWriter fallback)
# ────────────────────────────────────────────────
def pdfs_merge_v2(files):
    writer = PdfWriter()
    processed_pages = 0
    try:
        for idx, f in enumerate(files):
            if not is_pdf_filename(f.filename):
                continue
            f.seek(0)
            pdf_bytes = f.read()
            pdf_stream = io.BytesIO(pdf_bytes)
            try:
                reader = PdfReader(pdf_stream)
                for page in reader.pages:
                    writer.add_page(page)
                    processed_pages += 1
                print(f"✅ PDF {idx+1}: {f.filename} ({len(reader.pages)} páginas)")
            except Exception as e:
                print(f"❌ Error leyendo PDF {f.filename}: {e}")
                continue
        if processed_pages == 0:
            raise ValueError("No se pudieron procesar archivos PDF válidos.")
        output = io.BytesIO()
        writer.write(output)
        output.seek(0)
        return ("single", ("pdfs_unidos.pdf", output.read(), "application/pdf"))
    finally:
        try:
            writer.close()
        except:
            pass

# ────────────────────────────────────────────────
# Helpers para construir el DOCX
# ────────────────────────────────────────────────
def _add_table_border(table):
    """Agrega bordes visibles a una tabla DOCX."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.append(tblPr)

def _is_heading_line(line_text: str, font_size: float, page_avg_size: float) -> int:
    """
    Detecta si una línea es un encabezado y devuelve su nivel (1-3) o 0 si no lo es.
    Se basa en tamaño de fuente relativo y patrones de texto.
    """
    text = line_text.strip()
    if not text:
        return 0
    # Heading nivel 1: fuente bastante más grande que el promedio
    if font_size and page_avg_size:
        ratio = font_size / page_avg_size if page_avg_size > 0 else 1
        if ratio >= 1.5 and len(text) < 120:
            return 1
        if ratio >= 1.2 and len(text) < 120:
            return 2
        if ratio >= 1.05 and len(text) < 120:
            return 3
    # Patrones típicos de headings: todo mayúsculas corto
    if text.isupper() and 3 < len(text) < 80:
        return 2
    return 0

def _build_paragraph(doc, text: str, heading_level: int = 0,
                      bold: bool = False, font_size: float = None,
                      align: str = None):
    """Agrega un párrafo al documento con formato básico."""
    if heading_level > 0:
        p = doc.add_heading(text, level=min(heading_level, 4))
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if font_size:
            run.font.size = Pt(min(max(font_size, 6), 36))
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p

def _get_line_avg_size(line):
    """Obtiene el tamaño de fuente promedio de los chars de una línea."""
    sizes = [c.get('size') for c in line.get('chars', []) if c.get('size')]
    if not sizes:
        return None
    return sum(sizes) / len(sizes)

def _get_line_is_bold(line):
    """Detecta si la mayoría de chars de una línea son bold."""
    chars = line.get('chars', [])
    if not chars:
        return False
    bold_count = sum(1 for c in chars if 'Bold' in (c.get('fontname') or ''))
    return bold_count > len(chars) * 0.5

def _detect_alignment(line, page_width):
    """Detecta alineación aproximada de una línea."""
    x0 = line.get('x0', 0)
    x1 = line.get('x1', page_width)
    center = (x0 + x1) / 2
    page_center = page_width / 2
    if abs(center - page_center) < page_width * 0.05:
        return 'center'
    if x0 > page_width * 0.6:
        return 'right'
    return 'left'

# ────────────────────────────────────────────────
# PDF → DOCX  (con agrupación de párrafos)
# ────────────────────────────────────────────────

def _group_lines_into_paragraphs(lines, page_avg_size, page_width):
    """
    Agrupa líneas de pdfplumber en párrafos lógicos.

    Una línea pertenece al párrafo anterior cuando:
      - El gap vertical es <= 1.6 × line_height  (interlineado normal)
      - El x0 es similar (±15px) al x0 del párrafo (misma columna/indentación)

    Devuelve lista de dicts:
      {text, is_bold, font_size, align, is_heading_candidate}
    """
    if not lines:
        return []

    paragraphs = []
    current_parts = []       # lista de líneas del párrafo en construcción
    current_x0 = None
    prev_bottom = None
    line_height = page_avg_size * 1.4  # altura de línea estimada

    def flush():
        if not current_parts:
            return
        # Unir las líneas con espacio (texto justificado se corta en el margen)
        full_text = ' '.join(l['text'].strip() for l in current_parts if l['text'].strip())
        if not full_text:
            return
        # Tomar propiedades de la primera línea del grupo
        first = current_parts[0]
        avg_size = sum(_get_line_avg_size(l) or page_avg_size for l in current_parts) / len(current_parts)
        bold = _get_line_is_bold(first)
        align = _detect_alignment(first, page_width)
        heading_lvl = _is_heading_line(full_text, avg_size, page_avg_size)
        paragraphs.append({
            'text': full_text,
            'is_bold': bold,
            'font_size': avg_size,
            'align': align,
            'heading_lvl': heading_lvl,
        })

    for line in lines:
        text = line.get('text', '').strip()
        if not text:
            continue

        top = line.get('top', 0)
        x0 = line.get('x0', 0)

        gap = (top - prev_bottom) if prev_bottom is not None else 0
        same_indent = current_x0 is not None and abs(x0 - current_x0) < 15
        normal_spacing = gap <= line_height * 1.6

        if prev_bottom is None or (same_indent and normal_spacing):
            # Heurística extra: si la línea anterior terminó con punto o es muy corta
            # y esta nueva línea empieza desde el mismo margen, puede ser párrafo nuevo
            ends_block = False
            if current_parts:
                last_text = current_parts[-1].get('text', '').strip()
                ends_block = (
                    last_text.endswith('.')
                    or last_text.endswith(':')
                    or len(last_text) < 30
                )
            if ends_block and same_indent and gap > page_avg_size * 0.8:
                flush()
                current_parts = [line]
                current_x0 = x0
            else:
                current_parts.append(line)
                if current_x0 is None:
                    current_x0 = x0
        else:
            # Nuevo párrafo
            flush()
            current_parts = [line]
            current_x0 = x0

        prev_bottom = line.get('bottom', top + page_avg_size)

    flush()
    return paragraphs


def pdf_to_docx(files):
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    temp_dirs = []
    any_content = False

    # ── Deduplicar archivos por nombre (por si el frontend envía duplicados) ──
    seen_names = set()
    unique_files = []
    for f in files:
        name = f.filename.strip().lower()
        if name not in seen_names:
            seen_names.add(name)
            unique_files.append(f)
    files = unique_files

    try:
        for file_idx, f in enumerate(files):
            if not is_pdf_filename(f.filename):
                continue

            temp_dir = tempfile.mkdtemp()
            temp_path = os.path.join(temp_dir, secure_filename(f.filename))
            f.save(temp_path)
            temp_dirs.append(temp_dir)

            if file_idx > 0:
                doc.add_page_break()

            if len(files) > 1:
                doc.add_heading(f.filename, level=1)

            try:
                with pdfplumber.open(temp_path) as pdf:
                    total_pages = len(pdf.pages)

                    for page_num, page in enumerate(pdf.pages, 1):
                        page_width = page.width or 595

                        # Tamaño de fuente promedio
                        all_chars = page.chars or []
                        char_sizes = [c.get('size', 0) for c in all_chars if c.get('size', 0) > 0]
                        page_avg_size = (sum(char_sizes) / len(char_sizes)) if char_sizes else 11.0

                        # Separador de página (solo si hay más de una)
                        if total_pages > 1 and page_num > 1:
                            doc.add_page_break()

                        # ── Tablas ──
                        tables_bbox = []
                        try:
                            found_tables = page.find_tables()
                            for tbl_obj in found_tables:
                                tables_bbox.append(tbl_obj.bbox)
                                table_data = tbl_obj.extract()
                                if not table_data:
                                    continue
                                table_data = [row for row in table_data if any(
                                    cell and str(cell).strip() for cell in row
                                )]
                                if not table_data:
                                    continue
                                cols = max(len(row) for row in table_data)
                                rows_count = len(table_data)
                                if cols < 1 or rows_count < 1:
                                    continue
                                word_table = doc.add_table(rows=rows_count, cols=cols)
                                _add_table_border(word_table)
                                word_table.style = 'Table Grid'
                                for r_idx, row in enumerate(table_data):
                                    for c_idx, cell_text in enumerate(row):
                                        if c_idx >= cols:
                                            break
                                        word_table.cell(r_idx, c_idx).text = str(cell_text or '').strip()
                                doc.add_paragraph()
                                any_content = True
                        except Exception as te:
                            print(f"  ⚠️  Tablas p.{page_num}: {te}")

                        # ── Líneas de texto ──
                        try:
                            text_lines = page.extract_text_lines(layout=True, strip_whitespace=True)
                        except TypeError:
                            text_lines = page.extract_text_lines()

                        if not text_lines:
                            raw = page.extract_text()
                            if raw and raw.strip():
                                doc.add_paragraph(raw.strip())
                                any_content = True
                            continue

                        # Filtrar líneas que caen dentro de tablas
                        filtered_lines = []
                        for line in text_lines:
                            lx0 = line.get('x0', 0)
                            ly0 = line.get('top', 0)
                            lx1 = line.get('x1', page_width)
                            ly1 = line.get('bottom', ly0 + 10)
                            in_table = any(
                                bx0 <= lx0 and ly0 >= by0 and lx1 <= bx1 and ly1 <= by1
                                for bx0, by0, bx1, by1 in tables_bbox
                            )
                            if not in_table:
                                filtered_lines.append(line)

                        # Agrupar en párrafos lógicos
                        paragraphs = _group_lines_into_paragraphs(
                            filtered_lines, page_avg_size, page_width
                        )

                        prev_para_top = None
                        for para in paragraphs:
                            # Espacio entre bloques separados
                            _build_paragraph(
                                doc,
                                para['text'],
                                heading_level=para['heading_lvl'],
                                bold=para['is_bold'] and para['heading_lvl'] == 0,
                                font_size=para['font_size'] if para['heading_lvl'] == 0 else None,
                                align=para['align'] if para['heading_lvl'] == 0 else None,
                            )
                            any_content = True

            except Exception as e:
                print(f"❌ Error procesando {f.filename}: {e}")
                doc.add_paragraph(f"[Error al procesar {f.filename}: {e}]")

        if not any_content:
            doc.add_paragraph(
                "No se pudo extraer contenido de los PDFs.\n"
                "Es posible que sean PDFs escaneados (solo imágenes).\n"
                "Para ese caso se necesita OCR (por ejemplo Tesseract)."
            )

        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        if len(files) == 1:
            base_name = files[0].filename.rsplit(".", 1)[0]
            filename = f"{base_name}_convertido.docx"
        else:
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"pdfs_convertidos_{ts}.docx"

        return (
            "single",
            (filename, doc_buffer.read(),
             "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        )

    finally:
        for temp_dir in temp_dirs:
            try:
                shutil.rmtree(temp_dir)
            except:
                pass


# ────────────────────────────────────────────────
# Rutas Flask
# ────────────────────────────────────────────────
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        conversion = request.form.get("conversion")
        uploaded_files = request.files.getlist("files")
        valid_files = [f for f in uploaded_files
                       if f and f.filename and allowed_filename(f.filename)]
        if not valid_files:
            flash("No se subió ningún archivo válido.", "error")
            return redirect(url_for("index"))
        try:
            if conversion == "images_to_pdf":
                imgs = [f for f in valid_files if is_image_filename(f.filename)]
                if not imgs:
                    flash("No se encontraron imágenes válidas.", "error")
                    return redirect(url_for("index"))
                result_type, payload = images_to_pdf(imgs)

            elif conversion == "merge_pdfs":
                pdfs = [f for f in valid_files if is_pdf_filename(f.filename)]
                if len(pdfs) < 2:
                    flash("Se requieren al menos 2 archivos PDF para unir.", "error")
                    return redirect(url_for("index"))
                try:
                    result_type, payload = pdfs_merge_v1(pdfs)
                except Exception as e1:
                    print(f"❌ PdfMerger falló: {e1}")
                    result_type, payload = pdfs_merge_v2(pdfs)

            elif conversion == "pdf_to_docx":
                pdfs = [f for f in valid_files if is_pdf_filename(f.filename)]
                if not pdfs:
                    flash("No se encontraron archivos PDF.", "error")
                    return redirect(url_for("index"))
                result_type, payload = pdf_to_docx(pdfs)

            else:
                flash("Tipo de conversión no válido.", "error")
                return redirect(url_for("index"))

            if result_type == "single":
                filename, data, mime = payload
                response = make_response(data)
                response.headers.set("Content-Type", mime)
                response.headers.set("Content-Disposition", "attachment", filename=filename)
                return response

            flash("Error en la conversión.", "error")
            return redirect(url_for("index"))

        except Exception as e:
            flash(f"Error: {str(e)}", "error")
            return redirect(url_for("index"))

    return render_template('index.html')


@app.route("/health")
def health():
    return "OK", 200


# Ejecución directa (sin waitress)
if __name__ == "__main__":
    host = "0.0.0.0"
    port = 5050
    try:
        hostname = socket.gethostname()
        local_ip = socket.gethostbyname(hostname)
    except Exception:
        local_ip = "127.0.0.1"
    os.system('clear' if os.name == 'posix' else 'cls')
    print("=" * 60)
    print("XONI-CONVER v3.3 - Conversor Universal")
    print("=" * 60)
    print(f"URL Local:  http://{local_ip}:{port}")
    print(f"Móvil:      Usa la misma IP en tu red WiFi")
    print("=" * 60)
    url_completa = f"http://{local_ip}:{port}"
    print("\nESCANEA ESTE CÓDIGO QR PARA ACCEDER DESDE TU MÓVIL:\n")
    print(generar_qr_terminal(url_completa))
    print("\n" + "=" * 60)
    app.run(host=host, port=port, debug=False, threaded=True)
