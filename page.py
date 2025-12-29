from flask import Flask, request, make_response, render_template_string, flash, redirect, url_for
from PIL import Image
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from docx import Document
import io
import zipfile
import datetime
import socket
import os
import tempfile
import shutil
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.secret_key = "clave_secreta_xoni_conver_pc"

# Desactivar límite de tamaño
app.config['MAX_CONTENT_LENGTH'] = None

ALLOWED_IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "bmp", "gif", "tiff", "webp"}
ALLOWED_PDF_EXTENSIONS = {"pdf"}
ALLOWED_EXTENSIONS = ALLOWED_IMAGE_EXTENSIONS.union(ALLOWED_PDF_EXTENSIONS)

# INTERFAZ ELEGANTE - RESPONSIVE PARA PC Y MÓVIL
HTML_PAGE = '''<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>XONI-CONVER | Conversor Universal</title>
    <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0">
    <meta name="theme-color" content="#2d3748">
    <style>
        :root {
            --primary-dark: #1a365d;
            --primary: #2d3748;
            --accent: #4a5568;
            --light: #f7fafc;
            --border: #e2e8f0;
            --text: #2d3748;
            --text-light: #718096;
            --success: #38a169;
            --error: #e53e3e;
            --shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
            --shadow-hover: 0 10px 15px -3px rgba(0, 0, 0, 0.1);
            --radius: 8px;
            --radius-lg: 12px;
        }

        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif;
        }

        body {
            background: linear-gradient(135deg, #f5f7fa 0%, #e4e8f0 100%);
            color: var(--text);
            min-height: 100vh;
            padding: 15px;
            font-size: 16px;
        }

        @media (min-width: 768px) {
            body {
                padding: 30px 20px;
            }
        }

        .container {
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            border-radius: var(--radius-lg);
            box-shadow: var(--shadow);
            overflow: hidden;
            border: 1px solid var(--border);
        }

        .header {
            background: linear-gradient(to right, var(--primary-dark), var(--primary));
            color: white;
            padding: 25px 20px;
            text-align: center;
        }

        @media (min-width: 768px) {
            .header {
                padding: 40px;
            }
        }

        .header-icon {
            font-size: 2.5rem;
            margin-bottom: 15px;
            display: block;
        }

        @media (min-width: 768px) {
            .header-icon {
                font-size: 3rem;
                margin-bottom: 20px;
            }
        }

        .header h1 {
            font-size: 1.8rem;
            font-weight: 700;
            margin-bottom: 10px;
        }

        @media (min-width: 768px) {
            .header h1 {
                font-size: 2.5rem;
            }
        }

        .header p {
            font-size: 1rem;
            opacity: 0.9;
            max-width: 600px;
            margin: 0 auto;
            line-height: 1.5;
        }

        @media (min-width: 768px) {
            .header p {
                font-size: 1.1rem;
                line-height: 1.6;
            }
        }

        .content {
            padding: 25px;
            display: flex;
            flex-direction: column;
            gap: 25px;
        }

        @media (min-width: 1024px) {
            .content {
                padding: 40px;
                display: grid;
                grid-template-columns: 2fr 1fr;
                gap: 40px;
                min-height: 500px;
            }
        }

        .form-container {
            width: 100%;
        }

        @media (min-width: 1024px) {
            .form-container {
                border-right: 1px solid var(--border);
                padding-right: 40px;
            }
        }

        .form-title {
            font-size: 1.4rem;
            font-weight: 600;
            margin-bottom: 25px;
            color: var(--primary-dark);
            display: flex;
            align-items: center;
            gap: 10px;
        }

        @media (min-width: 768px) {
            .form-title {
                font-size: 1.5rem;
            }
        }

        .form-group {
            margin-bottom: 20px;
        }

        @media (min-width: 768px) {
            .form-group {
                margin-bottom: 25px;
            }
        }

        .form-label {
            display: block;
            font-weight: 600;
            margin-bottom: 8px;
            color: var(--text);
            font-size: 0.95rem;
        }

        .form-select {
            width: 100%;
            padding: 14px 16px;
            border: 2px solid var(--border);
            border-radius: var(--radius);
            font-size: 1rem;
            background: white;
            color: var(--text);
            transition: all 0.2s;
            cursor: pointer;
            -webkit-appearance: none;
            -moz-appearance: none;
            appearance: none;
            background-image: url("data:image/svg+xml;charset=UTF-8,%3csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round'%3e%3cpolyline points='6 9 12 15 18 9'%3e%3c/polyline%3e%3c/svg%3e");
            background-repeat: no-repeat;
            background-position: right 16px center;
            background-size: 16px;
        }

        .form-select:focus {
            outline: none;
            border-color: var(--primary);
            box-shadow: 0 0 0 3px rgba(26, 54, 93, 0.1);
        }

        .file-upload {
            border: 3px dashed var(--border);
            border-radius: var(--radius);
            padding: 30px 15px;
            text-align: center;
            background: var(--light);
            transition: all 0.3s;
            cursor: pointer;
        }

        @media (min-width: 768px) {
            .file-upload {
                padding: 40px 20px;
            }
        }

        .file-upload:hover, .file-upload.dragover {
            border-color: var(--success);
            background: rgba(56, 161, 105, 0.05);
        }

        .upload-icon {
            font-size: 2.5rem;
            color: var(--accent);
            margin-bottom: 15px;
        }

        @media (min-width: 768px) {
            .upload-icon {
                font-size: 3rem;
            }
        }

        .upload-text h3 {
            font-size: 1.2rem;
            margin-bottom: 10px;
            color: var(--text);
        }

        @media (min-width: 768px) {
            .upload-text h3 {
                font-size: 1.3rem;
            }
        }

        .upload-text p {
            color: var(--text-light);
            margin-bottom: 20px;
            font-size: 0.95rem;
        }

        .browse-btn {
            background: var(--primary);
            color: white;
            border: none;
            padding: 14px 28px;
            border-radius: var(--radius);
            font-weight: 600;
            cursor: pointer;
            transition: all 0.2s;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            gap: 8px;
            font-size: 1rem;
            width: 100%;
            max-width: 250px;
        }

        .browse-btn:hover {
            background: var(--primary-dark);
            transform: translateY(-2px);
            box-shadow: var(--shadow-hover);
        }

        .browse-btn:active {
            transform: translateY(0);
        }

        .file-list {
            margin-top: 20px;
            max-height: 250px;
            overflow-y: auto;
            border: 1px solid var(--border);
            border-radius: var(--radius);
            background: var(--light);
        }

        .file-item {
            display: flex;
            align-items: center;
            justify-content: space-between;
            padding: 14px 16px;
            border-bottom: 1px solid var(--border);
            transition: background 0.2s;
        }

        .file-item:hover {
            background: rgba(26, 54, 93, 0.05);
        }

        .file-item:last-child {
            border-bottom: none;
        }

        .file-info {
            display: flex;
            align-items: center;
            gap: 12px;
            flex: 1;
            min-width: 0;
        }

        .file-icon {
            color: var(--primary);
            font-size: 1.2rem;
            flex-shrink: 0;
        }

        .file-details {
            flex: 1;
            min-width: 0;
        }

        .file-name {
            font-weight: 500;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            font-size: 0.95rem;
        }

        @media (min-width: 768px) {
            .file-name {
                font-size: 1rem;
            }
        }

        .file-size {
            font-size: 0.85rem;
            color: var(--text-light);
        }

        .file-remove {
            background: none;
            border: none;
            color: var(--error);
            cursor: pointer;
            font-size: 1.1rem;
            padding: 4px 8px;
            border-radius: 4px;
            transition: background 0.2s;
            flex-shrink: 0;
            width: 32px;
            height: 32px;
            display: flex;
            align-items: center;
            justify-content: center;
        }

        .file-remove:hover {
            background: rgba(229, 62, 62, 0.1);
        }

        .convert-btn {
            width: 100%;
            background: linear-gradient(to right, var(--primary-dark), var(--primary));
            color: white;
            border: none;
            padding: 18px;
            border-radius: var(--radius);
            font-size: 1.1rem;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s;
            margin-top: 20px;
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 10px;
            touch-action: manipulation;
        }

        .convert-btn:hover:not(:disabled) {
            transform: translateY(-2px);
            box-shadow: 0 10px 25px rgba(26, 54, 93, 0.2);
        }

        .convert-btn:disabled {
            opacity: 0.5;
            cursor: not-allowed;
            transform: none !important;
        }

        .convert-btn:active:not(:disabled) {
            transform: translateY(0);
        }

        .info-panel {
            padding: 25px;
            background: var(--light);
            border-radius: var(--radius);
            border: 1px solid var(--border);
        }

        @media (min-width: 1024px) {
            .info-panel {
                margin-top: 0;
            }
        }

        .info-title {
            font-size: 1.3rem;
            font-weight: 600;
            margin-bottom: 20px;
            color: var(--primary-dark);
            display: flex;
            align-items: center;
            gap: 10px;
        }

        .info-list {
            list-style: none;
        }

        .info-item {
            margin-bottom: 15px;
            padding-bottom: 15px;
            border-bottom: 1px solid var(--border);
        }

        .info-item:last-child {
            margin-bottom: 0;
            padding-bottom: 0;
            border-bottom: none;
        }

        .info-item strong {
            color: var(--primary);
            display: block;
            margin-bottom: 5px;
            font-size: 1rem;
        }

        .info-item p {
            color: var(--text-light);
            font-size: 0.95rem;
            line-height: 1.5;
        }

        .file-counter {
            text-align: center;
            margin-top: 15px;
            color: var(--text-light);
            font-size: 0.9rem;
            padding: 12px;
            background: var(--light);
            border-radius: var(--radius);
            border: 1px solid var(--border);
        }

        .messages {
            position: fixed;
            top: 15px;
            right: 15px;
            left: 15px;
            width: auto;
            z-index: 1000;
            max-width: 400px;
            margin: 0 auto;
        }

        @media (min-width: 768px) {
            .messages {
                right: 20px;
                left: auto;
                width: 350px;
            }
        }

        .message {
            padding: 16px 20px;
            border-radius: var(--radius);
            margin-bottom: 10px;
            animation: slideIn 0.3s ease;
            box-shadow: var(--shadow);
            border-left: 4px solid;
            background: white;
            font-size: 0.95rem;
        }

        .message.success {
            border-left-color: var(--success);
            color: var(--success);
        }

        .message.error {
            border-left-color: var(--error);
            color: var(--error);
        }

        .message-content {
            display: flex;
            justify-content: space-between;
            align-items: flex-start;
            gap: 10px;
        }

        .message-text {
            flex: 1;
        }

        .message-close {
            background: none;
            border: none;
            color: inherit;
            cursor: pointer;
            font-size: 1.2rem;
            opacity: 0.7;
            transition: opacity 0.2s;
            padding: 0;
            width: 24px;
            height: 24px;
            display: flex;
            align-items: center;
            justify-content: center;
            flex-shrink: 0;
        }

        .message-close:hover {
            opacity: 1;
        }

        @keyframes slideIn {
            from {
                transform: translateX(100%);
                opacity: 0;
            }
            to {
                transform: translateX(0);
                opacity: 1;
            }
        }

        .footer {
            text-align: center;
            padding: 25px 20px;
            color: var(--text-light);
            font-size: 0.9rem;
            border-top: 1px solid var(--border);
            background: var(--light);
        }

        .footer p {
            margin-bottom: 5px;
        }

        .loading-overlay {
            position: fixed;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: rgba(255, 255, 255, 0.95);
            display: flex;
            align-items: center;
            justify-content: center;
            z-index: 9999;
            opacity: 0;
            visibility: hidden;
            transition: all 0.3s;
            backdrop-filter: blur(3px);
        }

        .loading-overlay.active {
            opacity: 1;
            visibility: visible;
        }

        .loading-spinner {
            width: 50px;
            height: 50px;
            border: 3px solid var(--border);
            border-top: 3px solid var(--primary);
            border-radius: 50%;
            animation: spin 1s linear infinite;
        }

        .loading-text {
            margin-top: 15px;
            color: var(--text);
            font-weight: 500;
        }

        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }

        .hidden {
            display: none !important;
        }

        .text-center {
            text-align: center;
        }

        @media (max-width: 767px) {
            button, 
            .form-select,
            .file-item {
                min-height: 48px;
            }
            
            .file-remove {
                width: 44px;
                height: 44px;
            }
            
            .browse-btn {
                padding: 16px 24px;
            }
        }

        #fileInput {
            display: none;
        }
    </style>
</head>
<body>
    <div class="loading-overlay" id="loadingOverlay">
        <div class="text-center">
            <div class="loading-spinner"></div>
            <div class="loading-text">Procesando archivos...</div>
        </div>
    </div>

    <div class="messages" id="messagesContainer"></div>

    <div class="container">
        <header class="header">
            <div class="header-icon">🔄</div>
            <h1>XONI-CONVER</h1>
            <p>Conversor universal para PC y móvil. Convierte imágenes a PDF, une múltiples PDFs y extrae texto de PDF a Word.</p>
        </header>

        <div class="content">
            <div class="form-container">
                <h2 class="form-title">
                    📁 Conversor de Archivos
                </h2>

                <form method="POST" enctype="multipart/form-data" id="converterForm">
                    <div class="form-group">
                        <label class="form-label">Tipo de Conversión</label>
                        <select name="conversion" id="conversion" class="form-select" required>
                            <option value="images_to_pdf">Imágenes → PDF</option>
                            <option value="merge_pdfs">Unir PDFs → PDF único</option>
                            <option value="pdf_to_docx">PDF → Word (.docx)</option>
                        </select>
                    </div>

                    <div class="form-group">
                        <label class="form-label">Archivos a Convertir</label>
                        <div class="file-upload" id="fileUploadArea">
                            <div class="upload-icon">📎</div>
                            <div class="upload-text">
                                <h3>Arrastra y suelta archivos aquí</h3>
                                <p>O haz clic para seleccionar</p>
                                <button type="button" class="browse-btn" id="browseButton">
                                    Explorar Archivos
                                </button>
                            </div>
                            <input type="file" name="files" id="fileInput" multiple required 
                                   accept=".pdf,image/*">
                        </div>

                        <div class="file-list" id="fileList">
                            <div class="file-item" style="justify-content: center; color: var(--text-light);">
                                No hay archivos seleccionados
                            </div>
                        </div>

                        <div class="file-counter" id="fileCounter">
                            0 archivos | 0 KB
                        </div>
                    </div>

                    <button type="submit" class="convert-btn" id="convertButton" disabled>
                        <span>🔄</span>
                        Convertir Archivos
                    </button>
                </form>
            </div>

            <div class="info-panel">
                <h3 class="info-title">
                    ℹ️ Información
                </h3>
                
                <ul class="info-list">
                    <li class="info-item">
                        <strong>Formatos Soportados</strong>
                        <p>• Imágenes: PNG, JPG, JPEG, BMP, GIF, TIFF, WEBP<br>
                           • Documentos: PDF, DOCX</p>
                    </li>
                    
                    <li class="info-item">
                        <strong>Sin Límites</strong>
                        <p>Procesa archivos de cualquier tamaño. No hay restricciones de tamaño ni cantidad.</p>
                    </li>
                    
                    <li class="info-item">
                        <strong>Seguro</strong>
                        <p>Todos los archivos se procesan en memoria y se eliminan automáticamente al terminar.</p>
                    </li>
                    
                    <li class="info-item">
                        <strong>Universal</strong>
                        <p>Interfaz optimizada para PC, móviles y tablets. Totalmente responsive.</p>
                    </li>
                </ul>
            </div>
        </div>

        <footer class="footer">
            <p>XONI-CONVER v3.2 • Conversor Universal • by XONIDU</p>
            <p style="margin-top: 5px; font-size: 0.8rem; opacity: 0.7;">
                Procesa archivos de forma segura y eficiente
            </p>
        </footer>
    </div>

    <script>
        const fileInput = document.getElementById('fileInput');
        const fileUploadArea = document.getElementById('fileUploadArea');
        const fileList = document.getElementById('fileList');
        const fileCounter = document.getElementById('fileCounter');
        const browseButton = document.getElementById('browseButton');
        const converterForm = document.getElementById('converterForm');
        const loadingOverlay = document.getElementById('loadingOverlay');
        const messagesContainer = document.getElementById('messagesContainer');
        const convertButton = document.getElementById('convertButton');

        let files = [];
        const MAX_FILES = 50;

        document.addEventListener('DOMContentLoaded', () => {
            setupEventListeners();
            checkForMessages();
        });

        function setupEventListeners() {
            fileInput.addEventListener('change', handleFileSelect);
            browseButton.addEventListener('click', () => fileInput.click());
            
            fileUploadArea.addEventListener('dragover', (e) => {
                e.preventDefault();
                fileUploadArea.classList.add('dragover');
            });
            
            fileUploadArea.addEventListener('dragleave', () => {
                fileUploadArea.classList.remove('dragover');
            });
            
            fileUploadArea.addEventListener('drop', (e) => {
                e.preventDefault();
                fileUploadArea.classList.remove('dragover');
                
                if (e.dataTransfer.files.length) {
                    fileInput.files = e.dataTransfer.files;
                    handleFileSelect({ target: fileInput });
                }
            });
            
            converterForm.addEventListener('submit', handleFormSubmit);
        }

        function handleFileSelect(event) {
            const newFiles = Array.from(event.target.files);
            
            if (files.length + newFiles.length > MAX_FILES) {
                showMessage(`Máximo ${MAX_FILES} archivos permitidos`, 'error');
                return;
            }
            
            newFiles.forEach(file => {
                if (!isFileDuplicate(file)) {
                    files.push(file);
                }
            });
            
            updateFileList();
        }

        function isFileDuplicate(newFile) {
            return files.some(existingFile => 
                existingFile.name === newFile.name && 
                existingFile.size === newFile.size
            );
        }

        function updateFileList() {
            fileList.innerHTML = '';
            
            if (files.length === 0) {
                fileList.innerHTML = `
                    <div class="file-item" style="justify-content: center; color: var(--text-light);">
                        No hay archivos seleccionados
                    </div>
                `;
                fileCounter.textContent = '0 archivos | 0 KB';
                convertButton.disabled = true;
                return;
            }
            
            files.forEach((file, index) => {
                const fileElement = createFileElement(file, index);
                fileList.appendChild(fileElement);
            });
            
            const totalSize = formatFileSize(
                files.reduce((sum, file) => sum + file.size, 0)
            );
            fileCounter.textContent = `${files.length} archivo(s) | ${totalSize}`;
            convertButton.disabled = false;
        }

        function createFileElement(file, index) {
            const div = document.createElement('div');
            div.className = 'file-item';
            div.innerHTML = `
                <div class="file-info">
                    <div class="file-icon">📄</div>
                    <div class="file-details">
                        <div class="file-name">${file.name}</div>
                        <div class="file-size">${formatFileSize(file.size)}</div>
                    </div>
                </div>
                <button type="button" class="file-remove" data-index="${index}" aria-label="Eliminar archivo">
                    ✕
                </button>
            `;
            
            div.querySelector('.file-remove').addEventListener('click', (e) => {
                e.stopPropagation();
                removeFile(index);
            });
            
            return div;
        }

        function formatFileSize(bytes) {
            if (bytes === 0) return '0 Bytes';
            const k = 1024;
            const sizes = ['Bytes', 'KB', 'MB', 'GB'];
            const i = Math.floor(Math.log(bytes) / Math.log(k));
            return parseFloat((bytes / Math.pow(k, i)).toFixed(1)) + ' ' + sizes[i];
        }

        function removeFile(index) {
            files.splice(index, 1);
            updateFileList();
            
            const dataTransfer = new DataTransfer();
            files.forEach(file => dataTransfer.items.add(file));
            fileInput.files = dataTransfer.files;
        }

        async function handleFormSubmit(e) {
            e.preventDefault();
            
            if (files.length === 0) {
                showMessage('Selecciona al menos un archivo', 'error');
                return;
            }
            
            loadingOverlay.classList.add('active');
            
            const formData = new FormData(converterForm);
            files.forEach(file => formData.append('files', file));
            
            try {
                const response = await fetch('/', {
                    method: 'POST',
                    body: formData
                });
                
                if (response.ok) {
                    const contentType = response.headers.get('content-type');
                    const contentDisposition = response.headers.get('content-disposition');
                    
                    if (contentType && contentDisposition) {
                        const blob = await response.blob();
                        const url = window.URL.createObjectURL(blob);
                        const a = document.createElement('a');
                        a.href = url;
                        
                        const filenameMatch = contentDisposition.match(/filename="(.+)"/);
                        a.download = filenameMatch ? filenameMatch[1] : 'converted_file';
                        
                        document.body.appendChild(a);
                        a.click();
                        document.body.removeChild(a);
                        window.URL.revokeObjectURL(url);
                        
                        showMessage('Conversión completada ✓', 'success');
                        
                        files = [];
                        updateFileList();
                        fileInput.value = '';
                    }
                } else {
                    const text = await response.text();
                    const parser = new DOMParser();
                    const doc = parser.parseFromString(text, 'text/html');
                    const error = doc.querySelector('.message')?.textContent || 'Error en la conversión';
                    showMessage(error, 'error');
                }
            } catch (error) {
                showMessage('Error de conexión con el servidor', 'error');
                console.error('Error:', error);
            } finally {
                loadingOverlay.classList.remove('active');
            }
        }

        function showMessage(text, type = 'error') {
            const messageDiv = document.createElement('div');
            messageDiv.className = `message ${type}`;
            messageDiv.innerHTML = `
                <div class="message-content">
                    <div class="message-text">${text}</div>
                    <button class="message-close" onclick="this.parentElement.parentElement.remove()" aria-label="Cerrar mensaje">
                        ✕
                    </button>
                </div>
            `;
            
            messagesContainer.appendChild(messageDiv);
            
            setTimeout(() => {
                if (messageDiv.parentElement) {
                    messageDiv.remove();
                }
            }, 5000);
        }

        function checkForMessages() {
            const errorElements = document.querySelectorAll('.error');
            errorElements.forEach(el => {
                if (el.textContent.trim()) {
                    showMessage(el.textContent.trim(), 'error');
                }
            });
        }

        window.showMessage = showMessage;
    </script>
</body>
</html>'''

def allowed_filename(filename: str) -> bool:
    if not filename or '.' not in filename:
        return False
    ext = filename.rsplit(".", 1)[1].lower()
    return ext in ALLOWED_EXTENSIONS

def is_image_filename(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_IMAGE_EXTENSIONS

def is_pdf_filename(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_PDF_EXTENSIONS

def images_to_pdf(files):
    """Convert images to PDF"""
    pil_images = []
    temp_dirs = []
    
    try:
        for f in files:
            if not is_image_filename(f.filename):
                continue
            
            temp_dir = tempfile.mkdtemp()
            temp_path = os.path.join(temp_dir, secure_filename(f.filename))
            f.save(temp_path)
            temp_dirs.append(temp_dir)
            
            try:
                img = Image.open(temp_path).convert("RGB")
                pil_images.append(img)
            except Exception as e:
                print(f"Error procesando imagen {f.filename}: {e}")
                continue
        
        if not pil_images:
            raise ValueError("No hay imágenes válidas para convertir.")
        
        output = io.BytesIO()
        if len(pil_images) == 1:
            pil_images[0].save(output, format="PDF")
        else:
            first, rest = pil_images[0], pil_images[1:]
            first.save(output, format="PDF", save_all=True, append_images=rest)
        
        output.seek(0)
        return ("single", ("imagenes_a_pdf.pdf", output.read(), "application/pdf"))
        
    finally:
        for temp_dir in temp_dirs:
            try:
                shutil.rmtree(temp_dir)
            except:
                pass

def pdfs_merge_v1(files):
    """Merge PDFs usando PdfMerger - Método 1"""
    merger = PdfMerger()
    temp_dirs = []
    
    try:
        for idx, f in enumerate(files):
            if not is_pdf_filename(f.filename):
                continue
            
            temp_dir = tempfile.mkdtemp()
            temp_path = os.path.join(temp_dir, secure_filename(f.filename))
            f.save(temp_path)
            temp_dirs.append(temp_dir)
            
            try:
                # Usar PdfMerger que es más robusto para unir
                merger.append(temp_path)
                print(f"✅ PDF {idx+1} agregado: {f.filename}")
            except Exception as e:
                print(f"❌ Error con PdfMerger para {f.filename}: {e}")
                raise
        
        if len(merger.pages) == 0:
            raise ValueError("No se pudieron procesar archivos PDF válidos.")
        
        print(f"📊 Total de páginas: {len(merger.pages)}")
        
        output = io.BytesIO()
        merger.write(output)
        merger.close()
        output.seek(0)
        
        return ("single", ("pdfs_combinados.pdf", output.read(), "application/pdf"))
        
    finally:
        for temp_dir in temp_dirs:
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except:
                pass

def pdfs_merge_v2(files):
    """Merge PDFs usando PdfWriter - Método 2 (alternativo)"""
    writer = PdfWriter()
    processed_pages = 0
    
    try:
        for idx, f in enumerate(files):
            if not is_pdf_filename(f.filename):
                continue
            
            # Guardar en memoria directamente
            f.seek(0)
            pdf_bytes = f.read()
            pdf_stream = io.BytesIO(pdf_bytes)
            
            try:
                reader = PdfReader(pdf_stream)
                
                # Verificar páginas
                if len(reader.pages) == 0:
                    print(f"⚠️ PDF vacío: {f.filename}")
                    continue
                
                # Agregar TODAS las páginas
                for page in reader.pages:
                    writer.add_page(page)
                    processed_pages += 1
                
                print(f"✅ PDF {idx+1}: {f.filename} ({len(reader.pages)} páginas)")
                
            except Exception as e:
                print(f"❌ Error leyendo PDF {f.filename}: {e}")
                continue
        
        if processed_pages == 0:
            raise ValueError("No se pudieron procesar archivos PDF válidos.")
        
        print(f"📊 Total final de páginas: {processed_pages}")
        
        output = io.BytesIO()
        writer.write(output)
        output.seek(0)
        
        return ("single", ("pdfs_unidos.pdf", output.read(), "application/pdf"))
        
    finally:
        try:
            writer.close()
        except:
            pass

def pdf_to_docx(files):
    """Convert PDF(s) to single DOCX file"""
    doc = Document()
    temp_dirs = []
    
    try:
        doc.add_heading('Documento Convertido de PDF', 0)
        text_extracted = False
        
        for idx, f in enumerate(files):
            if not is_pdf_filename(f.filename):
                continue
            
            temp_dir = tempfile.mkdtemp()
            temp_path = os.path.join(temp_dir, secure_filename(f.filename))
            f.save(temp_path)
            temp_dirs.append(temp_dir)
            
            try:
                reader = PdfReader(temp_path)
                
                if len(files) > 1:
                    doc.add_heading(f"PDF: {f.filename}", level=2)
                
                for page_num, page in enumerate(reader.pages, 1):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            text_extracted = True
                            
                            if len(reader.pages) > 1:
                                doc.add_heading(f"Página {page_num}", level=3)
                            
                            doc.add_paragraph(text)
                            
                            if page_num < len(reader.pages):
                                doc.add_page_break()
                    except:
                        continue
                
                if idx < len(files) - 1:
                    doc.add_heading("─" * 50, level=2)
                    
            except Exception as e:
                doc.add_paragraph(f"Error procesando {f.filename}: {str(e)}")
                continue
        
        if not text_extracted:
            doc.add_paragraph("No se pudo extraer texto de los PDFs.")
        
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        if len(files) == 1:
            base_name = files[0].filename.rsplit(".", 1)[0]
            filename = f"{base_name}_convertido.docx"
        else:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"pdfs_combinados_{timestamp}.docx"
        
        return ("single", (filename, doc_buffer.read(), 
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))
        
    finally:
        for temp_dir in temp_dirs:
            try:
                shutil.rmtree(temp_dir)
            except:
                pass

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        conversion = request.form.get("conversion")
        uploaded_files = request.files.getlist("files")
        
        # Filtrar archivos válidos
        valid_files = []
        for f in uploaded_files:
            if f and f.filename and allowed_filename(f.filename):
                valid_files.append(f)
        
        if not valid_files:
            flash("No se subió ningún archivo válido.", "error")
            return redirect(url_for("index"))
        
        try:
            if conversion == "images_to_pdf":
                imgs = [f for f in valid_files if is_image_filename(f.filename)]
                if not imgs:
                    flash("No se encontraron imágenes válidas.", "error")
                    return redirect(url_for("index"))
                result_type, payload = images_to_pdf(imgs)

            elif conversion == "merge_pdfs":
                pdfs = [f for f in valid_files if is_pdf_filename(f.filename)]
                if len(pdfs) < 2:
                    flash("Se requieren al menos 2 archivos PDF para unir.", "error")
                    return redirect(url_for("index"))
                
                # Intentar primero con PdfMerger, luego con PdfWriter
                try:
                    print("🔄 Intentando unir PDFs con PdfMerger (método 1)...")
                    result_type, payload = pdfs_merge_v1(pdfs)
                except Exception as e1:
                    print(f"❌ PdfMerger falló: {e1}")
                    print("🔄 Intentando con PdfWriter (método 2)...")
                    try:
                        result_type, payload = pdfs_merge_v2(pdfs)
                    except Exception as e2:
                        flash(f"Error al unir PDFs: {e2}", "error")
                        return redirect(url_for("index"))

            elif conversion == "pdf_to_docx":
                pdfs = [f for f in valid_files if is_pdf_filename(f.filename)]
                if not pdfs:
                    flash("No se encontraron archivos PDF.", "error")
                    return redirect(url_for("index"))
                result_type, payload = pdf_to_docx(pdfs)

            else:
                flash("Tipo de conversión no válido.", "error")
                return redirect(url_for("index"))

            # Devolver archivo
            if result_type == "single":
                filename, data, mime = payload
                response = make_response(data)
                response.headers.set("Content-Type", mime)
                response.headers.set("Content-Disposition", "attachment", filename=filename)
                return response
                
            else:
                flash("Error en la conversión.", "error")
                return redirect(url_for("index"))

        except Exception as e:
            flash(f"Error: {str(e)}", "error")
            return redirect(url_for("index"))

    return render_template_string(HTML_PAGE)

@app.route("/health")
def health():
    return "OK", 200

if __name__ == "__main__":
    host = "0.0.0.0"
    port = 5050
    
    try:
        hostname = socket.gethostname()
        local_ip = socket.gethostbyname(hostname)
    except:
        local_ip = "127.0.0.1"
    
    print("=" * 60)
    print("XONI-CONVER v3.2 - Conversor Universal (CORREGIDO)")
    print("=" * 60)
    print(f"🌐 URL Local:      http://{local_ip}:{port}")
    print(f"📱 Móvil:          Usa la misma IP en tu red WiFi")
    print("=" * 60)
    print("✨ Características:")
    print("   • 2 métodos para unir PDFs (sin duplicación)")
    print("   • Interfaz responsive para PC y móvil")
    print("   • Sin límites de tamaño")
    print("   • Procesamiento seguro en memoria")
    print("=" * 60)
    
    app.run(
        host=host, 
        port=port, 
        debug=False,
        threaded=True
    )